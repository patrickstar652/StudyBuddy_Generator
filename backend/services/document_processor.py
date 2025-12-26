"""
文件處理服務
處理 PDF、DOCX 等文件的解析和文字提取
"""

import os
from typing import List, Dict, Any
from PyPDF2 import PdfReader
from docx import Document
import tiktoken


class DocumentProcessor:
    def __init__(self):
        self.encoding = tiktoken.get_encoding("cl100k_base")
        self.chunk_size = 1000  # tokens per chunk
        self.chunk_overlap = 200  # overlap tokens
    
    def extract_text(self, file_path: str) -> str:
        """
        從文件中提取文字
        
        Args:
            file_path: 文件路徑
        
        Returns:
            提取的文字內容
        """
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        elif ext == '.txt':
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """從 PDF 提取文字"""
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
        return text.strip()
    
    def _extract_from_docx(self, file_path: str) -> str:
        """從 DOCX 提取文字"""
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # 也提取表格中的文字
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
        return text.strip()
    
    def _extract_from_txt(self, file_path: str) -> str:
        """從 TXT 提取文字"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # 嘗試其他編碼
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    def count_tokens(self, text: str) -> int:
        """計算文字的 token 數量"""
        return len(self.encoding.encode(text))
    
    def split_into_chunks(self, text: str) -> List[Dict[str, Any]]:
        """
        將文字分割成區塊用於嵌入
        
        Args:
            text: 完整文字
        
        Returns:
            區塊列表，每個區塊包含 content 和 metadata
        """
        tokens = self.encoding.encode(text)
        chunks = []
        
        start = 0
        chunk_index = 0
        
        while start < len(tokens):
            end = start + self.chunk_size
            
            # 取得區塊的 tokens
            chunk_tokens = tokens[start:end]
            
            # 解碼回文字
            chunk_text = self.encoding.decode(chunk_tokens)
            
            chunks.append({
                "content": chunk_text,
                "chunk_index": chunk_index,
                "token_count": len(chunk_tokens),
                "start_token": start,
                "end_token": min(end, len(tokens))
            })
            
            # 移動到下一個區塊，保留重疊
            start = end - self.chunk_overlap
            chunk_index += 1
        
        return chunks
    
    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """
        獲取文件資訊
        
        Args:
            file_path: 文件路徑
        
        Returns:
            文件資訊字典
        """
        text = self.extract_text(file_path)
        chunks = self.split_into_chunks(text)
        
        return {
            "total_characters": len(text),
            "total_tokens": self.count_tokens(text),
            "total_chunks": len(chunks),
            "file_size": os.path.getsize(file_path),
            "file_name": os.path.basename(file_path)
        }


def get_document_processor() -> DocumentProcessor:
    """獲取文件處理器實例"""
    return DocumentProcessor()
